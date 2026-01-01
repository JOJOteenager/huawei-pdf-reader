"""
文档处理器属性测试

Feature: huawei-pdf-reader
Property 1: 文档打开一致性
Property 2: 无效文档错误处理
Validates: Requirements 1.1, 1.2, 1.3, 1.6

测试文档处理器的核心功能属性。
"""

import sys
import tempfile
from pathlib import Path

# 添加 src 目录到 Python 路径
src_path = Path(__file__).parent.parent.parent / "src"
sys.path.insert(0, str(src_path))

import fitz  # PyMuPDF
from hypothesis import given, settings, strategies as st, assume
import pytest

from huawei_pdf_reader.document_processor import (
    PDFRenderer,
    WordRenderer,
    DocumentError,
    FileNotFoundError,
    UnsupportedFormatError,
    CorruptedFileError,
    create_renderer,
)
from huawei_pdf_reader.models import DocumentInfo


# ============== 辅助函数 ==============

def create_valid_pdf(path: Path, num_pages: int = 1, title: str = "") -> None:
    """创建有效的PDF文件用于测试"""
    doc = fitz.open()
    for i in range(num_pages):
        page = doc.new_page(width=595, height=842)  # A4尺寸
        page.insert_text((50, 50), f"Page {i + 1}", fontsize=12)
    
    if title:
        doc.set_metadata({"title": title})
    
    doc.save(str(path))
    doc.close()


def create_valid_docx(path: Path, paragraphs: list = None) -> None:
    """创建有效的Word文档用于测试"""
    from docx import Document
    
    doc = Document()
    if paragraphs:
        for para in paragraphs:
            doc.add_paragraph(para)
    else:
        doc.add_paragraph("Test paragraph")
    
    doc.save(str(path))


def create_corrupted_file(path: Path) -> None:
    """创建损坏的文件"""
    with open(path, 'wb') as f:
        f.write(b"This is not a valid PDF or DOCX file content")


# ============== 策略定义 ==============

# 有效页数策略 (1-20页)
valid_page_count_strategy = st.integers(min_value=1, max_value=20)

# 有效标题策略
valid_title_strategy = st.text(
    alphabet=st.characters(whitelist_categories=('L', 'N', 'P', 'S'), 
                           whitelist_characters=' '),
    min_size=0, 
    max_size=50
)

# 有效段落策略
valid_paragraph_strategy = st.text(
    alphabet=st.characters(whitelist_categories=('L', 'N', 'P', 'S'),
                           whitelist_characters=' '),
    min_size=1,
    max_size=200
)

# 段落列表策略
paragraphs_strategy = st.lists(valid_paragraph_strategy, min_size=1, max_size=10)


# ============== Property 1: 文档打开一致性 ==============

class TestDocumentOpenConsistency:
    """
    Property 1: 文档打开一致性
    
    For any 有效的PDF或Word文档路径，打开文档后返回的DocumentInfo应包含
    正确的文件类型、非零页数和有效的文档标题。
    
    Feature: huawei-pdf-reader, Property 1: 文档打开一致性
    Validates: Requirements 1.1, 1.2, 1.3
    """

    @given(
        num_pages=valid_page_count_strategy,
        title=valid_title_strategy
    )
    @settings(max_examples=100)
    def test_pdf_open_consistency(self, num_pages: int, title: str):
        """
        Property 1: PDF文档打开一致性
        
        For any 有效的PDF文档，打开后返回的DocumentInfo应包含：
        - 正确的文件类型 ('pdf')
        - 正确的页数 (等于创建时的页数)
        - 有效的标题 (非空字符串)
        
        Feature: huawei-pdf-reader, Property 1: 文档打开一致性
        Validates: Requirements 1.1, 1.3
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "test.pdf"
            create_valid_pdf(pdf_path, num_pages=num_pages, title=title)
            
            renderer = PDFRenderer()
            try:
                doc_info = renderer.open(pdf_path)
                
                # 验证文件类型
                assert doc_info.file_type == "pdf", \
                    f"Expected file_type 'pdf', got '{doc_info.file_type}'"
                
                # 验证页数非零且正确
                assert doc_info.total_pages > 0, \
                    f"Expected total_pages > 0, got {doc_info.total_pages}"
                assert doc_info.total_pages == num_pages, \
                    f"Expected {num_pages} pages, got {doc_info.total_pages}"
                
                # 验证标题有效（非空字符串）
                assert isinstance(doc_info.title, str), \
                    f"Expected title to be str, got {type(doc_info.title)}"
                assert len(doc_info.title) > 0, \
                    "Expected non-empty title"
                
                # 验证路径正确
                assert doc_info.path == pdf_path, \
                    f"Expected path {pdf_path}, got {doc_info.path}"
                
            finally:
                renderer.close()

    @given(paragraphs=paragraphs_strategy)
    @settings(max_examples=100, deadline=None)
    def test_word_open_consistency(self, paragraphs: list):
        """
        Property 1: Word文档打开一致性
        
        For any 有效的Word文档，打开后返回的DocumentInfo应包含：
        - 正确的文件类型 ('docx')
        - 非零页数
        - 有效的标题
        
        Feature: huawei-pdf-reader, Property 1: 文档打开一致性
        Validates: Requirements 1.2, 1.3
        """
        # 过滤掉只包含空白的段落
        valid_paragraphs = [p for p in paragraphs if p.strip()]
        assume(len(valid_paragraphs) > 0)
        
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "test.docx"
            create_valid_docx(docx_path, paragraphs=valid_paragraphs)
            
            renderer = WordRenderer()
            try:
                doc_info = renderer.open(docx_path)
                
                # 验证文件类型
                assert doc_info.file_type == "docx", \
                    f"Expected file_type 'docx', got '{doc_info.file_type}'"
                
                # 验证页数非零
                assert doc_info.total_pages > 0, \
                    f"Expected total_pages > 0, got {doc_info.total_pages}"
                
                # 验证标题有效
                assert isinstance(doc_info.title, str), \
                    f"Expected title to be str, got {type(doc_info.title)}"
                assert len(doc_info.title) > 0, \
                    "Expected non-empty title"
                
                # 验证路径正确
                assert doc_info.path == docx_path, \
                    f"Expected path {docx_path}, got {doc_info.path}"
                
            finally:
                renderer.close()


# ============== Property 2: 无效文档错误处理 ==============

class TestInvalidDocumentErrorHandling:
    """
    Property 2: 无效文档错误处理
    
    For any 无效的文件路径或损坏的文档，Document_Processor应返回错误而不是崩溃，
    且错误信息应非空。
    
    Feature: huawei-pdf-reader, Property 2: 无效文档错误处理
    Validates: Requirements 1.6
    """

    @given(filename=st.text(min_size=1, max_size=50).filter(lambda x: x.strip() and '/' not in x and '\\' not in x))
    @settings(max_examples=100)
    def test_nonexistent_pdf_error(self, filename: str):
        """
        Property 2: 不存在的PDF文件错误处理
        
        For any 不存在的文件路径，PDFRenderer应抛出FileNotFoundError，
        且错误信息非空。
        
        Feature: huawei-pdf-reader, Property 2: 无效文档错误处理
        Validates: Requirements 1.6
        """
        # 确保文件名有效
        safe_filename = "".join(c for c in filename if c.isalnum() or c in '._-')
        assume(len(safe_filename) > 0)
        
        nonexistent_path = Path(f"/nonexistent_dir_12345/{safe_filename}.pdf")
        
        renderer = PDFRenderer()
        try:
            with pytest.raises(FileNotFoundError) as exc_info:
                renderer.open(nonexistent_path)
            
            # 验证错误信息非空
            error_message = str(exc_info.value)
            assert len(error_message) > 0, "Expected non-empty error message"
        finally:
            renderer.close()

    @given(filename=st.text(min_size=1, max_size=50).filter(lambda x: x.strip() and '/' not in x and '\\' not in x))
    @settings(max_examples=100)
    def test_nonexistent_word_error(self, filename: str):
        """
        Property 2: 不存在的Word文件错误处理
        
        For any 不存在的文件路径，WordRenderer应抛出FileNotFoundError，
        且错误信息非空。
        
        Feature: huawei-pdf-reader, Property 2: 无效文档错误处理
        Validates: Requirements 1.6
        """
        safe_filename = "".join(c for c in filename if c.isalnum() or c in '._-')
        assume(len(safe_filename) > 0)
        
        nonexistent_path = Path(f"/nonexistent_dir_12345/{safe_filename}.docx")
        
        renderer = WordRenderer()
        try:
            with pytest.raises(FileNotFoundError) as exc_info:
                renderer.open(nonexistent_path)
            
            error_message = str(exc_info.value)
            assert len(error_message) > 0, "Expected non-empty error message"
        finally:
            renderer.close()

    def test_corrupted_pdf_error(self):
        """
        Property 2: 损坏的PDF文件错误处理
        
        For any 损坏的PDF文件，PDFRenderer应抛出CorruptedFileError，
        且错误信息非空。
        
        Feature: huawei-pdf-reader, Property 2: 无效文档错误处理
        Validates: Requirements 1.6
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            corrupted_path = Path(temp_dir) / "corrupted.pdf"
            create_corrupted_file(corrupted_path)
            
            renderer = PDFRenderer()
            try:
                with pytest.raises(CorruptedFileError) as exc_info:
                    renderer.open(corrupted_path)
                
                error_message = str(exc_info.value)
                assert len(error_message) > 0, "Expected non-empty error message"
            finally:
                renderer.close()

    def test_corrupted_word_error(self):
        """
        Property 2: 损坏的Word文件错误处理
        
        For any 损坏的Word文件，WordRenderer应抛出CorruptedFileError，
        且错误信息非空。
        
        Feature: huawei-pdf-reader, Property 2: 无效文档错误处理
        Validates: Requirements 1.6
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            corrupted_path = Path(temp_dir) / "corrupted.docx"
            create_corrupted_file(corrupted_path)
            
            renderer = WordRenderer()
            try:
                with pytest.raises(CorruptedFileError) as exc_info:
                    renderer.open(corrupted_path)
                
                error_message = str(exc_info.value)
                assert len(error_message) > 0, "Expected non-empty error message"
            finally:
                renderer.close()

    @given(extension=st.sampled_from(['.txt', '.jpg', '.png', '.mp3', '.zip', '.exe']))
    @settings(max_examples=100)
    def test_unsupported_format_error(self, extension: str):
        """
        Property 2: 不支持的文件格式错误处理
        
        For any 不支持的文件格式，create_renderer应抛出UnsupportedFormatError，
        且错误信息非空。
        
        Feature: huawei-pdf-reader, Property 2: 无效文档错误处理
        Validates: Requirements 1.6
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            unsupported_path = Path(temp_dir) / f"test{extension}"
            unsupported_path.write_text("test content")
            
            with pytest.raises(UnsupportedFormatError) as exc_info:
                create_renderer(unsupported_path)
            
            error_message = str(exc_info.value)
            assert len(error_message) > 0, "Expected non-empty error message"
